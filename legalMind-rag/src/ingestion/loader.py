"""
loader.py — Multi-format document loader.
Supports PDF, DOCX, TXT, HTML, and URLs.
Returns a normalized list of Document objects.
"""
import asyncio
from dataclasses import dataclass, field
from pathlib import Path
from typing import AsyncIterator
from urllib.parse import urlparse

import aiofiles
import aiohttp
import pdfplumber
from docx import Document as DocxDocument

from src.utils.logging import get_logger

log = get_logger(__name__)


@dataclass
class RawDocument:
    """Normalized document output from any loader."""

    text: str
    source: str
    metadata: dict = field(default_factory=dict)

    def __post_init__(self) -> None:
        self.metadata.setdefault("source", self.source)
        self.metadata.setdefault("char_count", len(self.text))


class DocumentLoader:
    """Unified async document loader.

    Usage:
        loader = DocumentLoader()
        async for doc in loader.load_path("./data/raw"):
            process(doc)
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".html"}

    # Public API

    async def load_file(self, path: Path | str) -> RawDocument:
        """Load a single file and return a RawDocument."""
        path = Path(path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return await self._load_pdf(path)
        elif ext in (".docx", ".doc"):
            return await self._load_docx(path)
        elif ext in (".txt", ".md"):
            return await self._load_text(path)
        elif ext == ".html":
            return await self._load_html(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def load_path(
        self, path: Path | str, recursive: bool = True
    ) -> AsyncIterator[RawDocument]:
        """Yield RawDocuments from all supported files under a directory."""
        path = Path(path)
        if path.is_file():
            yield await self.load_file(path)
            return

        pattern = "**/*" if recursive else "*"
        files = [
            f
            for f in path.glob(pattern)
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]
        log.info(f"Found {len(files)} documents to load", path=str(path))

        # Load files concurrently in batches of 10
        batch_size = 10
        for i in range(0, len(files), batch_size):
            batch = files[i : i + batch_size]
            tasks = [self.load_file(f) for f in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            for result, file in zip(results, batch):
                if isinstance(result, Exception):
                    log.error(f"Failed to load {file}", error=str(result))
                else:
                    yield result

    async def load_url(self, url: str) -> RawDocument:
        """Load a web page and return a RawDocument."""
        parsed = urlparse(url)
        if not parsed.scheme:
            raise ValueError(f"Invalid URL: {url}")

        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                resp.raise_for_status()
                content_type = resp.headers.get("Content-Type", "")
                html = await resp.text()

        # Strip HTML tags (simple approach; use BeautifulSoup for production)
        import re
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()

        return RawDocument(
            text=text,
            source=url,
            metadata={"source_type": "url", "domain": parsed.netloc},
        )

    # Private loaders

    async def _load_pdf(self, path: Path) -> RawDocument:
        """Extract text from PDF using pdfplumber (table-aware)."""
        log.debug(f"Loading PDF: {path.name}")
        loop = asyncio.get_event_loop()

        def _extract() -> tuple[str, dict]:
            pages_text = []
            metadata = {}
            with pdfplumber.open(path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                metadata["pdf_metadata"] = pdf.metadata or {}
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    # Also extract tables as plain text
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = " | ".join(str(c) for c in row if c)
                            if row_text.strip():
                                page_text += "\n" + row_text
                    pages_text.append(f"[Page {i + 1}]\n{page_text}")
            return "\n\n".join(pages_text), metadata

        text, meta = await loop.run_in_executor(None, _extract)
        return RawDocument(
            text=text,
            source=str(path),
            metadata={
                "source_type": "pdf",
                "filename": path.name,
                "file_size_bytes": path.stat().st_size,
                **meta,
            },
        )

    async def _load_docx(self, path: Path) -> RawDocument:
        """Extract text from Word document."""
        log.debug(f"Loading DOCX: {path.name}")
        loop = asyncio.get_event_loop()

        def _extract() -> str:
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                    if row_text.strip():
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)

        text = await loop.run_in_executor(None, _extract)
        return RawDocument(
            text=text,
            source=str(path),
            metadata={
                "source_type": "docx",
                "filename": path.name,
                "file_size_bytes": path.stat().st_size,
            },
        )

    async def _load_text(self, path: Path) -> RawDocument:
        """Load plain text or markdown file."""
        log.debug(f"Loading text: {path.name}")
        async with aiofiles.open(path, encoding="utf-8", errors="replace") as f:
            text = await f.read()
        return RawDocument(
            text=text,
            source=str(path),
            metadata={
                "source_type": path.suffix.lstrip("."),
                "filename": path.name,
                "file_size_bytes": path.stat().st_size,
            },
        )

    async def _load_html(self, path: Path) -> RawDocument:
        """Load and strip HTML file."""
        import re

        async with aiofiles.open(path, encoding="utf-8", errors="replace") as f:
            html = await f.read()
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()
        return RawDocument(
            text=text,
            source=str(path),
            metadata={"source_type": "html", "filename": path.name},
        )
